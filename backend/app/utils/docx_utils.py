from docx import Document


def extract_text_from_docx(file_path):
    doc = Document(file_path)
    full_text = []

    # Dapatkan indeks semua tabel dalam dokumen
    table_indices = {}
    for i, table in enumerate(doc.tables):
        # Cari parent paragraf dari tabel
        parent = table._element.getparent()
        while parent is not None and parent.tag != '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}body':
            parent = parent.getparent()
        if parent is not None:
            # Simpan indeks tabel dalam body
            index = list(parent).index(table._element)
            table_indices[index] = table

    # Iterasi melalui semua elemen dokumen
    current_table_index = 0
    for i, paragraph in enumerate(doc.paragraphs):
        # Cek apakah ada tabel sebelum paragraf ini
        while current_table_index in table_indices:
            table = table_indices[current_table_index]
            full_text.append("\n=== TABEL ===")
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                full_text.append("| " + " | ".join(cells) + " |")
            full_text.append("=== AKHIR TABEL ===\n")
            current_table_index += 1

        # Tambahkan paragraf jika tidak kosong
        if paragraph.text.strip():
            full_text.append(paragraph.text)

        current_table_index += 1

    # Cek jika masih ada tabel tersisa
    while current_table_index in table_indices:
        table = table_indices[current_table_index]
        full_text.append("\n=== TABEL ===")
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            full_text.append("| " + " | ".join(cells) + " |")
        full_text.append("=== AKHIR TABEL ===\n")
        current_table_index += 1

    return "\n".join(full_text)
